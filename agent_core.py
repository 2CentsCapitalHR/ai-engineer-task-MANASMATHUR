"""
agent_core.py
Core processing logic for the ADGM Corporate Agent.
This module:
- Detects document type from its text content.
- Checks if all required documents for a process are uploaded.
- Runs red flag detection on the content.
- Generates an annotated .docx file with review comments.
- Produces a JSON summary report.
"""

from pathlib import Path
from docx import Document
from utils.doc_parser import extract_paragraphs
from utils.checklist import DOC_TYPE_KEYWORDS, PROCESS_CHECKLISTS
from utils.red_flags import detect_red_flags_from_text
import json


def detect_doc_type(paragraphs):
    """
    Identify the type of document by matching keywords in its text.
    paragraphs: list of strings from the .docx file.
    returns: string name of the detected document type or "Unknown".
    """
    text = ' '.join(paragraphs).lower()
    scores = {}

    # Count how many keywords for each type appear in the document
    for doc_type, keywords in DOC_TYPE_KEYWORDS.items():
        match_count = sum(1 for kw in keywords if kw in text)
        scores[doc_type] = match_count

    # Pick the document type with the highest keyword match
    best_match = max(scores, key=scores.get)
    return best_match if scores[best_match] > 0 else "Unknown"


def verify_checklist(uploaded_doc_types):
    """
    Compare uploaded document types against the required ADGM checklist
    for company incorporation.
    """
    checklist = PROCESS_CHECKLISTS['company_incorporation']
    required_docs = checklist['required']

    missing = [doc for doc in required_docs if doc not in uploaded_doc_types]

    return {
        "process": checklist["display"],
        "documents_uploaded": len(uploaded_doc_types),
        "required_documents": len(required_docs),
        "missing_documents": missing
    }


def annotate_docx(input_path, output_path, findings):
    """
    Add review annotations at the end of a .docx file.
    findings: list of dictionaries containing issues found.
    """
    doc = Document(input_path)

    doc.add_paragraph('\n--- ADGM Automated Review Annotations ---\n')
    for issue in findings:
        doc.add_paragraph(
            f"ISSUE (Document: {issue.get('document', 'Unknown')}): "
            f"{issue.get('section', 'N/A')} - {issue['issue']}"
        )
        doc.add_paragraph(
            f"Severity: {issue['severity']}; Suggestion: {issue['suggestion']}"
        )

    doc.save(output_path)


def analyze_and_run(paths, out_dir):
    """
    Main pipeline:
    1. Read each uploaded document.
    2. Detect its type.
    3. Check for missing required documents.
    4. Detect legal red flags.
    5. Annotate the first document.
    6. Output a JSON summary and the annotated file.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    uploaded_types = []
    issues_found = []

    for file_path in paths:
        paragraphs = extract_paragraphs(file_path)
        doc_type = detect_doc_type(paragraphs)
        uploaded_types.append(doc_type)

        # Detect issues in the document text
        text_content = ' '.join(paragraphs)
        doc_issues = detect_red_flags_from_text(text_content)

        # Tag issues with the original file name
        for issue in doc_issues:
            issue['document'] = Path(file_path).name
        issues_found.extend(doc_issues)

    # Compare against checklist
    checklist_result = verify_checklist([t for t in uploaded_types if t != 'Unknown'])
    missing_doc = checklist_result['missing_documents'][0] if checklist_result['missing_documents'] else None

    # Prepare JSON output
    json_summary = {
        "process": checklist_result["process"],
        "documents_uploaded": checklist_result["documents_uploaded"],
        "required_documents": checklist_result["required_documents"],
        "missing_document": missing_doc,
        "issues_found": issues_found
    }

    # Annotate the first document (if available)
    annotated_path = None
    if paths:
        annotated_path = out_dir / f"annotated_{Path(paths[0]).name}"
        annotate_docx(paths[0], annotated_path, issues_found)

    # Save JSON to file
    report_path = out_dir / 'analysis.json'
    with open(report_path, 'w') as f:
        json.dump(json_summary, f, indent=2)

    return {
        "annotated_path": str(annotated_path) if annotated_path else None,
        "report_path": str(report_path),
        "json": json_summary
    }


if __name__ == '__main__':
    # For quick testing from the terminal
    example_docs = list(Path('examples').glob('*.docx'))
    if example_docs:
        result = analyze_and_run([str(example_docs[0])], 'examples')
        print(json.dumps(result['json'], indent=2))
    else:
        print("No example documents found.")
