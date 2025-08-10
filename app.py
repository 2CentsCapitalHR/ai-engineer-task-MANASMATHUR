import gradio as gr
import docx
from docx.shared import RGBColor
import os

def analyze(file):
    doc = docx.Document(file)
    annotated_doc = docx.Document()

    for para in doc.paragraphs:
        text = para.text
        if "emergency light" in text.lower():
            run = annotated_doc.add_paragraph().add_run(text)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Highlight in red
        else:
            annotated_doc.add_paragraph(text)

    output_path = "annotated.docx"
    annotated_doc.save(output_path)
    return output_path

with gr.Blocks() as demo:
    gr.Markdown("# Emergency Lighting Fixture Detector")
    
    with gr.Row():
        uploader = gr.File(label="Upload Blueprint (.docx)", type="filepath")
        download = gr.File(label="Annotated .docx (Download)", type="filepath")
    
    analyze_btn = gr.Button("Analyze")
    
    analyze_btn.click(
        fn=analyze,
        inputs=uploader,
        outputs=download
    )

if __name__ == "__main__":
    demo.launch()
