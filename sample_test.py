from docx import Document
from pathlib import Path

Path("examples").mkdir(exist_ok=True)

doc = Document()
doc.add_heading('Company Incorporation Document', 0)
doc.add_paragraph('This document includes all necessary details for company incorporation.')
doc.add_paragraph('Memorandum of Association is included here.')
doc.add_paragraph('Other sections with relevant keywords.')

doc.save('examples/sample1.docx')
print("Sample document saved as 'examples/sample1.docx'")
